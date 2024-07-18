from sheets import *

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_COLOR_INDEX


def add_head_paragraph(document, conference_name):  # Добавляем шапку документа
    p1 = document.add_paragraph()
    text = f"""Программа {conference_name}
    по кафедре № 43 компьютерных технологий и программной 
    инженерии"""
    run = p1.add_run(text)
    run.bold = True
    run.italic = True
    p1.alignment = 1
    run.add_break()

    p2 = document.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.25)
    text = "Секция каф.43. «компьютерных технологий и программной инженерии»"
    run = p2.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(12)
    p1.alignment = 1

    p3 = document.add_paragraph()
    p3.paragraph_format.left_indent = Cm(2)
    p3.add_run("Научный руководитель секции – Колесникова Светлана Ивановна\n").font.size = Pt(12)
    p3.add_run("проф., д.т.н., проф.\n").font.size = Pt(12)
    p3.add_run("Зам. научного руководителя секции – Попов Алексей Анатольевич\n").font.size = Pt(12)
    p3.add_run("доц., к.т.н.\n").font.size = Pt(12)
    p3.add_run("Секретарь – Поляк Марк Дмитриевич\n").font.size = Pt(12)
    p3.add_run("ст. преп.").font.size = Pt(12)


def add_session_paragraph(document):    # Добавляем текст с заседанием
    p = document.add_paragraph()
    p.add_run("Заседание 1\n").bold = True
    run = p.add_run("Дата и время\n")
    run.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
    run.font.size = Pt(12)


def set_document_base_style(document):  # меняем шрифт документа и отступы
    document_line_spacing = document.styles["Normal"]
    document_line_spacing.paragraph_format.space_after = Pt(0)
    document_font = document.styles["Normal"].font
    document_font.name = "Times New Roman"
    document_font.size = Pt(14)


def set_field_size(document):   # задаем поля и размеры страницы документа
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_article_list(document, values):  # добавляем список статей
    for item in values:
        buff1 = document.add_paragraph(style='List Number')
        buff1.add_run(f"{item.author.get_fullname()}, группа {item.author.group}")
        buff1.paragraph_format.left_indent = Cm(1.25)
        buff1.paragraph_format.space_after = Pt(0)
        buff2 = document.add_paragraph()
        buff2.add_run(f"{item.title}")
        buff2.paragraph_format.first_line_indent = Cm(1.25)
        buff2.paragraph_format.space_after = Pt(10)
        buff2.paragraph_format.alignment = 3


def get_conference_program(spreadsheet_id, conference_name):
    document = Document()

    set_document_base_style(document)  # меняем шрифт документа и отступы
    set_field_size(document)  # задаем поля и размер страниц документа

    add_head_paragraph(document, conference_name)   # добавляем шапку документа

    add_session_paragraph(document)  # добавляем текст о заседании

    creds = get_credentials()
    values = get_sheet_values(creds, spreadsheet_id)  # получаем данные для списка статей

    add_article_list(document, values)  # добавляем список статей

    if os.path.exists("Программа к43.docx"):  # если такой файл уже есть, то удаляем
        os.remove("Программа к43.docx")

    document.save('Программа к43.docx')