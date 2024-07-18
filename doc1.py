from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from decoration import add_paragraph_with_style
from decoration import add_numbered_paragraph

def generate_document_programme(reports, short_name):
    doc = Document()

    # Добавление заголовка
    add_paragraph_with_style(doc, f'Программа {short_name}\nпо кафедре № 43 компьютерных технологий и программной\nинженерии', font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_paragraph_with_style(doc, 'Секция каф.43. «компьютерных технологий и программной инженерии»', font_size=12, bold=True)
    add_paragraph_with_style(doc, 'Научный руководитель секции – Колесникова Светлана Ивановна проф., д.т.н., проф.', font_size=12, left_indent=0.5)
    add_paragraph_with_style(doc, 'Зам. научного руководителя секции – Попов Алексей Анатольевич доц., к.т.н.', font_size=12, left_indent=0.5)
    add_paragraph_with_style(doc, 'Секретарь – Поляк Марк Дмитриевич ст. преп.', font_size=12, left_indent=0.5, space_after=20)

    add_paragraph_with_style(doc, 'Заседание 1.', font_size=14, bold=True)
    add_paragraph_with_style(doc, '20 апреля, 09:00, ауд. 23-10 БМ.', font_size=12, bold=True, space_after=20)
    add_paragraph_with_style(doc, 'По решению руководителя секции порядок следования докладов может быть изменен.', font_size=14, space_after=20)

    # Заполнение документа
    for index, row in reports.iterrows():
        # ФИО автора
        authors = f"{row['surname']} {row['name']}"
        if row['patronymic']:
            authors += f" {row['patronymic']}"
        # ФИО соавторов
        if len(row['coauthors']) > 1:
            coauthors = row['coauthors']
            coauthors_text = ", ".join([
                    f"{coauthor['surname']} {coauthor['name']}" + (f" {coauthor['patronymic']}" if coauthor['patronymic'] else '')
                    for coauthor in coauthors
                    if coauthor['surname'] != row['surname']
            ])
            authors = f"{authors}, {coauthors_text}"

        # группа автора
        if row['student_group']:
            text = f"{authors}, группа {row['student_group']}\n{row['title']}"
        else:
            text = f"{authors}\n{row['title']}"

        add_numbered_paragraph(doc, text, font_size=14)

    return doc