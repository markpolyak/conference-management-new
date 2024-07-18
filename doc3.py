from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from decoration import add_paragraph_with_style
from decoration import add_numbered_paragraph

def generate_document_publications(reports, short_name):
    doc = Document()

    # Добавление заголовка
    add_paragraph_with_style(doc, 'Список представляемых к публикации докладов', font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_paragraph_with_style(doc, 'Кафедра № 43 компьютерных технологий и программной инженерии', font_size=12, left_indent=0.5)
    add_paragraph_with_style(doc, 'Поляк Марк Дмитриевич', font_size=12, left_indent=0.5)
    add_paragraph_with_style(doc, 'e-mail: m.polyak@guap.ru', font_size=12, left_indent=0.5)
    add_paragraph_with_style(doc, 'тел.: +7-…', font_size=12, left_indent=0.5, space_after=20)

    # Заполнение документа
    for index, row in reports.iterrows():
        if int(row['presentation_quality']) != 0:
            authors = f"{row['surname']} {row['name'][0]}."
            if row['patronymic']:
                authors += f"{row['patronymic'][0]}."

            if row['coauthors'] != []:
                coauthors = row['coauthors']
                coauthors_text = ", ".join([
                    f"{coauthor['surname']} {coauthor['name'][0]}." + (f"{coauthor['patronymic'][0]}." if coauthor['patronymic'] else '')
                    for coauthor in coauthors
                    if coauthor['surname'] != row['surname']
                ])
                authors = f"{authors}, {coauthors_text}"

            text = f"{authors} {row['title']}"
            add_numbered_paragraph(doc, text, font_size=12, font_name='Times New Roman')

    add_paragraph_with_style(doc, '\nРуководитель УНИДС\t\t\t\t\tПоляк М.Д.', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, left_indent=0.5)

    return doc