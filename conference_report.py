from sheets import *

import os

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_COLOR_INDEX


def add_head_paragraph(document, conference_name):  # Добавляем шапку документа
    p1 = document.add_paragraph()
    text = f"""Отчет о проведении {conference_name}
    Секция 43. Кафедра компьютерных технологий и программной инженерии"""
    run = p1.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    p1.alignment = 1


def add_session_paragraph(document):    # Добавляем текст с заседанием
    p = document.add_paragraph()
    p.add_run("Заседание 1\n").bold = True
    p.add_run("Дата и время").font.highlight_color = WD_COLOR_INDEX.GRAY_25
    p.add_run("				")
    p.add_run("Место").font.highlight_color = WD_COLOR_INDEX.GRAY_25
    p.add_run("\nНаучный руководитель секции – ")
    p.add_run("Руководитель").font.highlight_color = WD_COLOR_INDEX.GRAY_25
    p.add_run("\nСекретарь – ")
    p.add_run("Секретарь").font.highlight_color = WD_COLOR_INDEX.GRAY_25
    p.paragraph_format.space_after = Pt(0)


def add_table_paragraph(document, values):  # добавляем таблицу в документ
    p = document.add_paragraph("\nСписок докладов\n")
    p.paragraph_format.space_after = Pt(0)

    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№ п/п"
    hdr_cells[1].text = "Фамилия и инициалы докладчика, название доклада"
    hdr_cells[2].text = "Статус (магистр / студент)"
    hdr_cells[3].text = "Решение"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    iter = 0
    for item in values:
        row_cells = table.add_row().cells
        iter += 1
        row_cells[0].text = str(iter)
        row_cells[1].text = f"{item.author.get_fullname()}. {item.title}"

        if item.author.group == "":
            row_cells[2].text = ""
        elif item.author.group[-1] in ["М", "м"]:
            row_cells[2].text = f"Магистрант гр.{item.author.group}"
        else:
            row_cells[2].text = f"Студент гр.{item.author.group}"

        if item.presentation_quality == "0":
            row_cells[3].text = "Доклад плохо подготовлен"
        elif item.presentation_quality == "1":
            row_cells[3].text = "Опубликовать доклад в сборнике СНК"
        elif item.presentation_quality == "2":
            row_cells[3].text = "Опубликовать доклад в сборнике СНК. Рекомендовать к участию в финале конкурса на лучшую студенческую научную работу ГУАП "
        else:
            row_cells[3].text = ""

    widths = (Cm(1.06), Cm(8.82), Cm(2.47), Cm(3.53))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            row.cells[idx].paragraphs[0].left_indent = Cm(0.18)
            row.cells[idx].paragraphs[0].alignment = 1

    document.add_paragraph("").paragraph_format.space_after = Pt(0)


def set_document_base_styles(document):  # Меняем шрифт документа
    document_font = document.styles["Normal"].font
    document_font.name = "Times New Roman"
    document_font.size = Pt(10)


def set_field_size(document):   # Задаем поля и размер страницы
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_footer(document):   # Добавляем текст для подписи
    p1 = document.add_paragraph()
    text = "\nНаучный руководитель секции                                    _________________ / Охтилев М.Ю."
    run = p1.add_run(text)


def get_conference_report(spreadsheet_id, conference_name):  # создаем и сохраняем документ
    document = Document()

    set_document_base_styles(document)  # задаем шрифт документа
    set_field_size(document)  # задаем поля и размеры документа

    add_head_paragraph(document, conference_name)   # добавляем шапку документа

    creds = get_credentials()
    values = get_sheet_values(creds, spreadsheet_id)    # получаем данные для создания таблицы

    add_session_paragraph(document)  # добавляем текст о заседании
    add_table_paragraph(document, values)  # добавляем таблицу

    add_footer(document)  # добавляем текст для подписи

    if os.path.exists("Отчёт.docx"):    # если такой файл уже есть, то удаляем его
        os.remove("Отчёт.docx")

    document.save('Отчёт.docx')