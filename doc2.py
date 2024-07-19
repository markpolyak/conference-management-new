from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from decoration import add_paragraph_with_style
from decoration import set_font_size
from decoration import set_cell_border

def generate_document_report(reports, short_name):
    doc = Document()

    # Добавление заголовка
    add_paragraph_with_style(doc, f'Отчет о проведении {short_name}', font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_with_style(doc, 'Секция 43. Кафедра компьютерных технологий и программной инженерии', font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_paragraph_with_style(doc, 'Заседание 1.', font_size=10, bold=True)
    add_paragraph_with_style(doc, '20 апреля, 09-00 \t\t ул. Б. Морская, д. 67, ауд. 23-10', font_size=10)
    add_paragraph_with_style(doc, 'Научный руководитель секции – д-р техн. наук, проф. М.Ю. Охтилев', font_size=10)
    add_paragraph_with_style(doc, 'Секретарь – д-р техн. наук, проф. С.И. Колесникова', font_size=10, space_after=20)


    add_paragraph_with_style(doc, 'Список докладов', font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=20)

    # Создание таблицы с нужным количеством столбцов
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    headers = ['№ п/п', 'Фамилия и инициалы докладчика, название доклада', 'Статус (магистр/ студент)', 'Решение']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

        set_font_size(hdr_cells[i], 10, True)

        set_cell_border(hdr_cells[i], top={"sz": 12, "val": "single", "color": "000000"},
                                      bottom={"sz": 12, "val": "single", "color": "000000"},
                                      start={"sz": 12, "val": "single", "color": "000000"},
                                      end={"sz": 12, "val": "single", "color": "000000"})

    # Заполнение таблицы
    for index, row in reports.iterrows():
        row_cells = table.add_row().cells
        print(row)
        print(reports.columns)
        # Порядковый номер
        row_cells[0].text = str(index + 1)

        # ФИО + доклад
        text = f"{row['surname']} {row['name']}"
        if row['patronymic'] != '':
            text = text + f" {row['patronymic']}"

        # Соавторы

        coauthors = row['coauthors']
        # text = text + ' совместно с ' + ", ".join(f"{coauthor['Surname']} {coauthor['Name'][0]}." + (f"{coauthor['Patronymic'][0]}." if coauthor['Patronymic'] else "") + ' подготовили доклад на тему:' + f" \"{row['Tittle_report']}.\""
        # text = text + "," + ", ".join(
        #         f" {coauthor['Surname']} {coauthor['Name']}" + (f" {coauthor['Patronymic']}" if coauthor['Patronymic'] else "")
        #         for coauthor in coauthors
        #     ) + ' подготовили доклад на тему:' + f" \"{row['Tittle_report']}.\""
        if len(row['coauthors']) > 1:
            text = text + ", " + ", ".join([
                    f"{coauthor['surname']} {coauthor['name']}" + (f" {coauthor['patronymic']}" if coauthor['patronymic'] else '')
                    for coauthor in coauthors
                    if coauthor['surname'] != row['surname']
                ]) + ' подготовили доклад на тему:' + f" \"{row['title']}.\""
        else:
            text = text + '.' + f" {row['title']}."

        row_cells[1].text = text

        # Роль пользователя
        if row['applicant_role'].lower() == 'студент':
            status = 'Студент'
            if row['student_group'] != '':
                status = status + '\nгр. ' + row['student_group'] 
        elif row['applicant_role'].lower() == 'магистрант':
            status = 'Магистрант'
            if row['student_group'] != '':
                status = status + '\nгр. ' + row['student_group'] 
        elif row['applicant_role'].lower() == 'сотрудник':
            status = 'Сотрудник'

        # Статус доклада
        row_cells[2].text = status
        if row['presentation_quality'] == '0':
            analysis = "Доклад плохо подготовлен"
        elif row['presentation_quality'] == '1':
            analysis = "Опубликовать доклад в сборнике МСНК"
        elif row['presentation_quality'] == '2':
            analysis = "Опубликовать доклад в сборнике МСНК; рекомендовать к участию в финале конкурса на лучшую студенческую научную работу ГУАП"
        else:
            analysis = "Неизвестный статус"

        # Решение
        row_cells[3].text = analysis

        for cell in row_cells:
            set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"},
                                  bottom={"sz": 12, "val": "single", "color": "000000"},
                                  start={"sz": 12, "val": "single", "color": "000000"},
                                  end={"sz": 12, "val": "single", "color": "000000"})

            set_font_size(cell, 10, False)
    # Установка ширины первой колонки
    for cell in table.columns[0].cells:
        cell.width = Pt(10)
    table.columns[0].width = Pt(10)

    # Установка ширины третьей колонки
    for cell in table.columns[1].cells:
        cell.width = Pt(200)
    table.columns[1].width = Pt(200)

    add_paragraph_with_style(doc, '\n\nНаучный руководитель секции \t\t_________________ / Охтилев М.Ю.', font_size=10)

    return doc