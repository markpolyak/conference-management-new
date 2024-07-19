from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_paragraph_with_style(doc, text, font_size=12, bold=False, font_name='Times New Roman', alignment=WD_ALIGN_PARAGRAPH.LEFT, left_indent=None, first_line_indent=None, space_after=0):
    """
    Добавляет параграф с заданными стилями в документ.
    
    :param doc: Документ, в который будет добавлен параграф.
    :param text: Текст параграфа.
    :param font_size: Размер шрифта.
    :param bold: Жирный шрифт.
    :param font_name: Имя шрифта.
    :param alignment: Выравнивание параграфа.
    :param left_indent: Отступ слева (Inches).
    :param first_line_indent: Отступ первой строки (Inches).
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = font_name
    paragraph.alignment = alignment
    if left_indent is not None:
        paragraph.paragraph_format.left_indent = Inches(left_indent)
    if first_line_indent is not None:
        paragraph.paragraph_format.first_line_indent = Inches(first_line_indent)

    paragraph.paragraph_format.space_after = Pt(space_after)
    return paragraph

def add_numbered_paragraph(doc, text, font_size=12, font_name='Times New Roman', alignment=WD_ALIGN_PARAGRAPH.LEFT, left_indent=None):
    """
    Добавляет нумерованный параграф в документ.
    """
    paragraph = doc.add_paragraph(style='List Number')
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(0)
    if left_indent is not None:
        paragraph.paragraph_format.left_indent = Inches(left_indent)
    return paragraph

def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    for border_name in ["top", "start", "bottom", "end"]:
        border = kwargs.get(border_name)
        if border:
            element = OxmlElement(f"w:{border_name}")
            for key in ["sz", "val", "color", "space"]:
                if key in border:
                    element.set(qn(f"w:{key}"), str(border[key]))
            tcPr.append(element)

def set_font_size(cell, size, bold=False, name='Times New Roman'):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            rPr = run._element.get_or_add_rPr()
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(size * 2))
            rPr.append(sz)

            if bold:
                b = OxmlElement('w:b')
                rPr.append(b)

            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), name)
            rFonts.set(qn('w:hAnsi'), name)
            rPr.append(rFonts)