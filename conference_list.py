from sheets import *

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_COLOR_INDEX


def add_head_paragraph(document):   # добавляем шапку документа
    p1 = document.add_paragraph()
    run = p1.add_run("Список представляемых к публикации докладов\n")
    run.bold = True
    run.italic = True
    run.font.size = Pt(14)
    p1.alignment = 1

    p2 = document.add_paragraph()
    p2.paragraph_format.left_indent = Cm(2)
    p2.add_run("Кафедра № 43 компьютерных технологий и программной инженерии").font.size = Pt(12)
    name = p2.add_run("\nФИО")
    name.font.size = Pt(12)
    name.font.highlight_color = WD_COLOR_INDEX.GRAY_25
    p2.add_run("\ne-mail: ").font.size = Pt(12)
    email = p2.add_run("Почта")
    email.font.size = Pt(12)
    email.font.highlight_color = WD_COLOR_INDEX.GRAY_25
    p2.add_run("\nтел.: ").font.size = Pt(12)
    phone = p2.add_run("номер\n")
    phone.font.size = Pt(12)
    phone.font.highlight_color = WD_COLOR_INDEX.GRAY_25


def add_article_list(document, values):  # Добавляем список статей
    for item in values:
        buff1 = document.add_paragraph(style='List Number')
        buff1.add_run(f"{item.author.get_initials()} {item.title}")
        buff1.paragraph_format.first_line_indent = Cm(1.25)
        buff1.paragraph_format.space_after = Pt(0)


def set_field_size(document):   # задаем поля и размер страниц документа
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_document_base_style(document):  # меняем шрифт документа и отступы
    document_line_spacing = document.styles["Normal"]
    document_line_spacing.paragraph_format.space_after = Pt(0)
    document_font = document.styles["Normal"].font
    document_font.name = "Times New Roman"
    document_font.size = Pt(14)


def add_footer(document):
    p1 = document.add_paragraph()
    text = "\nРуководитель УНИДС                                               Поляк М.Д."
    run = p1.add_run(text)
    run.font.size = Pt(12)
    p1.paragraph_format.left_indent = Cm(2)


def get_conference_list(spreadsheet_id):
    document = Document()

    set_document_base_style(document)  # меняем шрифт документа и отступы
    set_field_size(document)  # задаем поля и размер страниц документа

    add_head_paragraph(document)  # добавляем шапку документа

    creds = get_credentials()
    values = get_sheet_values(creds, spreadsheet_id)    # получаем данные для создания списка статей

    add_article_list(document, values)  # добавляем список статей

    add_footer(document)  # добавляем текст для подписи

    if os.path.exists("Список.docx"):  # если такой файл уже есть, то удаляем
        os.remove("Список.docx")

    document.save("Список.docx")